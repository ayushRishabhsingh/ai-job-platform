"""Resume text extraction.

PDF -> text is done with PyMuPDF (fast, no external binaries). DOCX and plain
text are supported as a convenience. Everything degrades gracefully: if an
optional library is missing, the caller gets a clear error instead of a crash
at import time.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

SUPPORTED_CONTENT_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
    "text/markdown": "txt",
}

SUPPORTED_EXTENSIONS = {".pdf": "pdf", ".docx": "docx", ".txt": "txt", ".md": "txt"}


class UnsupportedResumeFormat(Exception):
    """Raised when the uploaded file is not a format we can read."""


@dataclass
class ParsedResume:
    text: str
    page_count: int = 0
    word_count: int = 0
    sections: dict[str, str] = field(default_factory=dict)


def detect_kind(filename: str, content_type: str | None) -> str:
    """Prefer the file extension, fall back to the declared MIME type."""
    ext = Path(filename).suffix.lower()
    if ext in SUPPORTED_EXTENSIONS:
        return SUPPORTED_EXTENSIONS[ext]
    if content_type and content_type in SUPPORTED_CONTENT_TYPES:
        return SUPPORTED_CONTENT_TYPES[content_type]
    raise UnsupportedResumeFormat(f"Cannot read '{filename}'. Supported formats: PDF, DOCX, TXT.")


def _extract_pdf(path: Path) -> tuple[str, int]:
    try:
        import pymupdf  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise UnsupportedResumeFormat(
            "PDF support requires PyMuPDF. Install it with: pip install pymupdf"
        ) from exc

    chunks: list[str] = []
    with pymupdf.open(path) as doc:
        for page in doc:
            chunks.append(page.get_text("text"))
        pages = doc.page_count
    return "\n".join(chunks), pages


def _extract_docx(path: Path) -> tuple[str, int]:
    try:
        import docx  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise UnsupportedResumeFormat(
            "DOCX support requires python-docx. Install it with: pip install python-docx"
        ) from exc

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts), 1


def _extract_txt(path: Path) -> tuple[str, int]:
    return path.read_text(encoding="utf-8", errors="ignore"), 1


def clean_text(text: str) -> str:
    """Normalise whitespace and the bullet/ligature noise PDFs love to emit."""
    text = text.replace("\u00a0", " ").replace("\ufb01", "fi").replace("\ufb02", "fl")
    text = re.sub(r"[•▪◦●·]", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


SECTION_PATTERNS = {
    "skills": r"(technical\s+skills|skills|technologies|tech\s+stack|competencies)",
    "experience": r"(work\s+experience|experience|employment|professional\s+experience)",
    "education": r"(education|academic|qualifications)",
    "projects": r"(projects|personal\s+projects|academic\s+projects)",
    "certifications": r"(certifications?|licenses?|courses?)",
}


def split_sections(text: str) -> dict[str, str]:
    """Best-effort split of a resume into named sections.

    Heuristic on purpose: resumes have no schema. Used to weight the SKILLS
    section more heavily during extraction.
    """
    lines = text.split("\n")
    hits: list[tuple[int, str]] = []
    for idx, line in enumerate(lines):
        stripped = line.strip().lower()
        if not stripped or len(stripped) > 60:
            continue
        for name, pattern in SECTION_PATTERNS.items():
            if re.fullmatch(rf"{pattern}\s*:?", stripped):
                hits.append((idx, name))
                break

    sections: dict[str, str] = {}
    for i, (start, name) in enumerate(hits):
        end = hits[i + 1][0] if i + 1 < len(hits) else len(lines)
        sections[name] = "\n".join(lines[start + 1 : end]).strip()
    return sections


def parse_resume(
    path: str | Path, filename: str | None = None, content_type: str | None = None
) -> ParsedResume:
    """Read a resume file from disk and return cleaned text + metadata."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Resume file not found: {path}")

    kind = detect_kind(filename or path.name, content_type)
    if kind == "pdf":
        raw, pages = _extract_pdf(path)
    elif kind == "docx":
        raw, pages = _extract_docx(path)
    else:
        raw, pages = _extract_txt(path)

    text = clean_text(raw)
    return ParsedResume(
        text=text,
        page_count=pages,
        word_count=len(text.split()),
        sections=split_sections(text),
    )
